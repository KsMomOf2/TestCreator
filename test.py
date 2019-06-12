import docx

#import openpyxl

path = 'C:\\My Documents\\Coding Club\\TestCreator\\'
file = 'SeniorEnglishFinal'

def getTest():
	doc = docx.Document(path + file + '.docx')
	c = doc.paragraphs[0].text
	t = doc.paragraphs[1].text
	i = doc.paragraphs[3].text

	qs = []
	for p in range(4,len(doc.paragraphs)):
		paragraph = doc.paragraphs[p]
		q = paragraph.text
		for r in paragraph.runs:
			if r.bold:
				print("Answer")
			print("run: " + r.style.name)
		qs.append(q)
		print("par: " + paragraph.style.name)

	return c, t, i, qs

# findAnswer takes in a list of questions with their mutliple choice answers
# and identifies the correct answer.  Puts the correct answer in the first 
# part of the multiple choice

def findAnswer(questions):
	for q in questions:
			print(q)
	return q

course, title, instructions, questions = getTest()

print (course + ' of ' + title)
print (instructions)
for q in questions:
	print(q)




# xls = openpyxl.load_workbook(path + file + '.xlsx')
#sheet = xls.active
#sheet.title = 'master'

# Read the Word File



# Grab Class, Exam, Instructions

# Loop through all questions
# For each questions
# save the question in Column 1
# save the answer (bolded) in Column 2
# save the rest of the answers in columns 3..n

# close the word file
#save the excel file

#sheet.create_sheet(title='newsheetname')
#doc.add_paragraph('Hello world!')
#doc.add_paragraph('First Item',style=doc.styles['List'])
#doc.add_paragraph('Second Item',style=doc.styles['List'])
#oc.add_paragraph('Third Item',style=doc.styles['List'])
#doc.save(path + file)
print("I'm done")